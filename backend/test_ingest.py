"""Raw-doc ingest self-checks (pytest).

Covers: .md Obsidian cleanup, .docx (python-docx), .pdf (pymupdf), and the
author-optional rule in main.create_job (doc-only needs no author; an email
file with no author is rejected with 400).

Run: cd backend && DATA_DIR=./.data uv run pytest -q
"""
from __future__ import annotations

import asyncio
import json
import os
import tempfile

# Isolate main's module-level DBOS launch into a throwaway data dir BEFORE the
# import (main reads DATA_DIR/DBOS_DB at import time).
_TMPDATA = tempfile.mkdtemp(prefix="sc-test-")
os.environ["DATA_DIR"] = _TMPDATA
os.environ.setdefault("DBOS_DB", f"sqlite:///{_TMPDATA}/dbos.sqlite")
os.environ.setdefault("OPENROUTER_API_KEY", "test-key")

import ingest  # noqa: E402
import main  # noqa: E402  (triggers DBOS.launch — harmless, just SQLite)


# ── ingest: .md cleanup ────────────────────────────────────────────────────
def test_md_cleanup(tmp_path):
    md = (
        "---\n"
        "title: My Note\n"
        "tags: [project]\n"
        "---\n"
        "This note references [[Alice|Alicia]] and [[Bob]] with #project/alpha.\n\n"
        "A [real link](https://example.com/page) sits here for the cleanup test.\n\n"
        "Body prose runs long enough to clear the minimum-char threshold so the\n"
        "sample is kept rather than dropped as too short and discarded outright.\n\n"
        "> [!note] Callout whose marker should vanish but whose text should remain.\n"
    )
    (tmp_path / "note.md").write_text(md)
    samples = ingest.ingest(tmp_path, set())
    assert samples, "expected at least one sample from the .md"
    joined = " ".join(s["text"] for s in samples)

    # frontmatter stripped
    assert "title: My Note" not in joined
    assert "tags: [project]" not in joined
    # wikilinks resolved (display/target kept, brackets gone)
    assert "[[" not in joined and "]]" not in joined
    assert "Alicia" in joined and "Bob" in joined
    # tag dropped
    assert "#project/alpha" not in joined
    # link: label kept, url gone
    assert "real link" in joined and "example.com" not in joined
    # callout marker dropped, line text kept
    assert "[!note]" not in joined
    assert "marker should vanish" in joined


# ── ingest: .docx ───────────────────────────────────────────────────────────
def test_docx_extracted(tmp_path):
    import docx  # python-docx

    document = docx.Document()
    document.add_paragraph("First paragraph of an essay about the weather today.")
    document.add_paragraph("")  # empty paragraph — must be skipped
    document.add_paragraph("Second paragraph keeps the thought going at length.")
    document.save(str(tmp_path / "essay.docx"))

    samples = ingest.ingest(tmp_path, set())
    assert samples, "expected at least one sample from the .docx"
    joined = " ".join(s["text"] for s in samples)
    assert "First paragraph of an essay" in joined
    assert "Second paragraph keeps the thought" in joined


# ── ingest: .pdf ────────────────────────────────────────────────────────────
def test_pdf_extracted(tmp_path):
    import fitz  # pymupdf

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text(
        (72, 72),
        "This sentence is deliberately long enough to clear the minimum "
        "sample character threshold in the PDF extraction path of ingest.")
    doc.save(str(tmp_path / "scan.pdf"))
    doc.close()

    samples = ingest.ingest(tmp_path, set())
    assert samples, "expected at least one sample from the .pdf"
    joined = " ".join(s["text"] for s in samples)
    assert "clear the minimum" in joined


# ── main.create_job: author-optional rule ───────────────────────────────────
class _StubFile:
    """Minimal UploadFile stand-in: create_job only touches filename + read()."""

    def __init__(self, name, data=b""):
        self.filename = name
        self._data = data

    async def read(self):
        return self._data


def test_doc_only_without_author_ok(tmp_path, monkeypatch):
    # A plain-text doc never needs an author; create_job must accept it.
    monkeypatch.setattr(main, "DATA", tmp_path)
    monkeypatch.setattr(main, "start_job", lambda *a, **k: None)
    monkeypatch.setattr(main.status, "create_job", lambda *a, **k: None)

    f = _StubFile("writing.txt", b"some prose that needs no author filter at all.")
    resp = asyncio.run(main.create_job(author="", files=[f]))
    assert resp.status_code == 200
    assert "job_id" in json.loads(resp.body)


def test_email_without_author_rejected(tmp_path, monkeypatch):
    # An email file with no author address can't be filtered → clear 400.
    monkeypatch.setattr(main, "DATA", tmp_path)
    monkeypatch.setattr(main, "start_job", lambda *a, **k: None)
    monkeypatch.setattr(main.status, "create_job", lambda *a, **k: None)

    f = _StubFile("inbox.mbox", b"From someone@example.com ...\n")
    resp = asyncio.run(main.create_job(author="", files=[f]))
    assert resp.status_code == 400
